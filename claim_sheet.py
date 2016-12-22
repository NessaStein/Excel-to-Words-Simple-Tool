# coding=gbk
from openpyxl import load_workbook
import dateutil.parser
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

document = Document()
wb = load_workbook(filename='ҩ�����׺�Ʒ��¼.xlsx', read_only=True)
ws = wb['Sheet1']

row_number = 0
table_row_index = 1
table_total_cost = 0.00

for row in ws.rows:
    row_number += 1
    if row_number < 3:
        continue
    if row[1].value is None:
        break

    if row[0].value is not None:
        if table_total_cost != 0.00:
            current_table.cell(1, 6).merge(current_table.cell(table_row_index, 6))
            current_table.cell(1, 7).merge(current_table.cell(table_row_index, 7))
            current_table.cell(1, 8).merge(current_table.cell(table_row_index, 8))

            row_cells = current_table.add_row().cells
            row_cells[0].merge(row_cells[8])
            row_cells[0].text = u"                                                                            �ϼ�" + str(
                table_total_cost) + u"Ԫ"
            paragraph = document.add_paragraph(" ")
            paragraph = document.add_paragraph(" ")
            paragraph = document.add_paragraph(" ")

        paragraph = document.add_paragraph(u"XXX΢�����о���ʵ��������õ�")
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        date_info = dateutil.parser.parse(str(row[0].value)).strftime('%Y-%m-%d')
        detail_info = u"ҩ�������                                                         ��Ʊ�ţ�" + str(row[7].value)

        paragraph = document.add_paragraph(detail_info)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        current_table = document.add_table(rows=1, cols=9)
        current_table.style = 'TableGrid'
        current_table.autofit = True
        current_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = current_table.rows[0].cells
        hdr_cells[0].text = u"���"
        hdr_cells[1].text = u"���Ƽ����"
        hdr_cells[2].text = u"��λ"
        hdr_cells[3].text = u"����"
        hdr_cells[4].text = u"���ۣ�Ԫ��"
        hdr_cells[5].text = u"��Ԫ��"
        hdr_cells[6].text = u"������"
        hdr_cells[7].text = u"����ʱ��"
        hdr_cells[8].text = u"֧����Ŀ"

        table_row_index = 1
        table_total_cost = 0.00
    else:
        table_row_index += 1

    table_total_cost += float(row[6].value)
    row_cells = current_table.add_row().cells
    # _Row()
    row_cells[0].text = str(table_row_index)
    for i in range(1, 6):
        cell_text = " "
        cell_value = row[i + 1].value
        if cell_value is not None:
            if isinstance(cell_value, unicode):
                cell_text = cell_value
            else:
                cell_text = str(cell_value)

        row_cells[i].text = cell_text

document.save('claim_sheet.docx')
